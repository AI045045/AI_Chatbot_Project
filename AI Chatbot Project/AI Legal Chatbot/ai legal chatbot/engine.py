"""
AI Legal Chatbot - RAG Engine

Manages the core LangChain RAG pipeline. This includes document ingestion (PDF, DOCX),
text splitting (RecursiveCharacterTextSplitter), local FAISS vector search, and local
Ollama embeddings and LLM generation (e.g. nomic-embed-text + llama3).
"""

import os
import re
import textwrap
import urllib.error
import urllib.request
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain_community.vectorstores import FAISS

from pypdf import PdfReader
from docx import Document as DocxDocument


def get_ollama_base_url() -> str:
    """Ollama HTTP API base (same env var the Ollama CLI uses)."""
    return os.environ.get("OLLAMA_HOST", "http://localhost:11434").rstrip("/")


def check_ollama_reachable(base_url: Optional[str] = None) -> Tuple[bool, str]:
    """
    Return (True, "") if the Ollama server responds, else (False, error detail).
    """
    base = (base_url or get_ollama_base_url()).rstrip("/")
    url = f"{base}/api/tags"
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "AI-LegalChatBot"})
        with urllib.request.urlopen(req, timeout=5) as resp:
            if resp.status == 200:
                return True, ""
    except urllib.error.HTTPError as e:
        return False, f"HTTP {e.code}: {e.reason}"
    except Exception as e:
        return False, str(e)
    return False, "Unexpected response from Ollama"


def ollama_setup_hint() -> str:
    return (
        "1) Install and start **Ollama** (or run `ollama serve` in a terminal). "
        "2) Pull models: `ollama pull nomic-embed-text` and `ollama pull llama3:8b` "
        "(or set the names in Model settings). "
        f"3) Default API URL: `{get_ollama_base_url()}` — override with env `OLLAMA_HOST` if needed."
    )


def _safe_filename(name: str) -> str:
    # Keep filenames filesystem-friendly.
    name = os.path.basename(name)
    name = re.sub(r"[^a-zA-Z0-9._-]+", "_", name).strip("._")
    return name or "upload"


def load_document(file_path: str, filename: str) -> List[Document]:
    """
    Load a PDF/DOCX document locally and return LangChain Documents.

    For PDFs, we create one Document per page to enable page-level citations.
    """
    filename = _safe_filename(filename)
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _load_pdf(file_path=file_path, filename=filename)
    if ext in {".docx", ".doc"}:
        return _load_docx(file_path=file_path, filename=filename)

    raise ValueError("Unsupported file type. Please upload a PDF or DOCX.")


def _load_pdf(file_path: str, filename: str) -> List[Document]:
    reader = PdfReader(file_path)
    docs: List[Document] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue

        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": filename,
                    "page": i + 1,
                },
            )
        )
    return docs


def _load_docx(file_path: str, filename: str) -> List[Document]:
    doc = DocxDocument(file_path)

    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Add basic table extraction if present (best-effort).
    # python-docx table iteration can get large; keep it simple.
    for table in doc.tables:
        for row in table.rows:
            cell_texts = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    cell_texts.append(ct)
            if cell_texts:
                parts.append(" | ".join(cell_texts))

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("DOCX loaded but contained no readable text.")

    return [
        Document(
            page_content=text,
            metadata={
                "source": filename,
                "page": None,
            },
        )
    ]


def split_into_chunks(
    docs: List[Document],
    *,
    chunk_size: int = 1200,
    chunk_overlap: int = 200,
) -> List[Document]:
    """
    Split documents into overlapping chunks for retrieval.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_documents(docs)

    for idx, ch in enumerate(chunks):
        src = ch.metadata.get("source", "upload")
        page = ch.metadata.get("page", None)
        page_part = f"p{page}" if page else "no-page"
        ch.metadata = dict(ch.metadata)
        ch.metadata["chunk_index"] = idx
        ch.metadata["chunk_id"] = f"{src}|{page_part}|c{idx}"

    return chunks


@dataclass
class RAGSession:
    """
    Holds the in-memory RAG state for a Streamlit session.

    Uses FAISS in RAM (no disk persistence). On Windows this avoids ChromaDB's
    native `chroma-hnswlib` build requirement.
    """

    vectorstore: Any
    llm_model: str
    embedding_model: str
    llm: ChatOllama

    def cleanup(self) -> None:
        """Drop references so the index and model handles can be garbage-collected."""
        self.vectorstore = None  # type: ignore
        self.llm = None  # type: ignore


def build_rag_session_from_files(
    *,
    file_path: str,
    filename: str,
    embedding_model: str,
    llm_model: str,
    chunk_size: int = 1200,
    chunk_overlap: int = 200,
) -> RAGSession:
    """
    Load, chunk, embed, and build an ephemeral vector store for a single document.
    """
    raw_docs = load_document(file_path=file_path, filename=filename)
    chunks = split_into_chunks(
        raw_docs, chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )

    base_url = get_ollama_base_url()
    embeddings = OllamaEmbeddings(model=embedding_model, base_url=base_url)
    vectorstore = FAISS.from_documents(chunks, embeddings)

    llm = ChatOllama(model=llm_model, temperature=0, base_url=base_url)

    return RAGSession(
        vectorstore=vectorstore,
        llm_model=llm_model,
        embedding_model=embedding_model,
        llm=llm,
    )


def _format_context(docs: List[Document], *, max_context_chars: int = 12000) -> str:
    """
    Produce a context block where each chunk has a stable citation id.
    """
    parts: List[str] = []
    total = 0

    for d in docs:
        chunk_id = d.metadata.get("chunk_id", "unknown-chunk")
        src = d.metadata.get("source", "unknown-source")
        page = d.metadata.get("page", None)
        page_part = f"page {page}" if page else "page unknown"
        chunk_text = d.page_content.strip()

        # Hard truncate individual chunks to keep prompts bounded.
        if len(chunk_text) > 4000:
            chunk_text = chunk_text[:4000] + "..."

        block = (
            f"[{chunk_id}] (source={src}, {page_part})\n"
            f"{chunk_text}\n"
        )
        if total + len(block) > max_context_chars:
            break
        parts.append(block)
        total += len(block)

    return "\n".join(parts)


def answer_question(
    session: RAGSession,
    question: str,
    *,
    top_k: int = 4,
    max_context_chars: int = 12000,
) -> Dict[str, Any]:
    """
    Privacy-first RAG:
    - Retrieves the most relevant chunks from the uploaded document only.
    - Instructs the LLM to answer ONLY using the provided context.
    - Returns answer text + citation metadata for UI display.
    """
    if not session.vectorstore:
        raise ValueError("RAG session not initialized. Upload and process a document first.")

    retrieved = session.vectorstore.similarity_search(question, k=top_k)
    context = _format_context(retrieved, max_context_chars=max_context_chars)

    prompt = textwrap.dedent(
        f"""
        You are a privacy-preserving legal assistant.
        Use ONLY the provided CONTEXT to answer the user's question.
        If the CONTEXT does not contain the answer, reply exactly:
        "I don't know based on the provided document."

        CONTEXT:
        {context}

        USER QUESTION:
        {question}

        Provide a concise answer. When possible, include the citation id(s) you used
        exactly as shown in square brackets (e.g., [legal.pdf|p2|c7]).
        """
    ).strip()

    result = session.llm.invoke(prompt)
    answer_text = getattr(result, "content", None) or str(result)

    citations = []
    for d in retrieved:
        citations.append(
            {
                "chunk_id": d.metadata.get("chunk_id"),
                "source": d.metadata.get("source"),
                "page": d.metadata.get("page"),
                "text_preview": d.page_content.strip()[:500],
            }
        )

    return {
        "answer": answer_text,
        "citations": citations,
    }

